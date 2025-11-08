"""
Haystack Pipeline Implementation
Handles document ingestion, indexing, retrieval, and generation using Haystack
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import json
import requests

# Document processing
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document as DocxDocument

# Haystack components
from haystack import Document, Pipeline
from haystack.components.writers import DocumentWriter
from haystack.components.preprocessors import DocumentCleaner, DocumentSplitter
from haystack.components.embedders import SentenceTransformersDocumentEmbedder, SentenceTransformersTextEmbedder
from haystack.components.retrievers import InMemoryBM25Retriever, InMemoryEmbeddingRetriever
from haystack.components.rankers import TransformersSimilarityRanker
from haystack.components.builders import PromptBuilder
from haystack.components.generators import HuggingFaceLocalGenerator
from haystack.document_stores.in_memory import InMemoryDocumentStore
from haystack.utils.device import ComponentDevice

# Transformers (only needed for local inference)
from transformers import AutoTokenizer, AutoModelForCausalLM
import torch

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes different document formats"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = extract_pdf_text(file_path)
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def process_document(file_path: str) -> Optional[Document]:
        """Process a document and return a Haystack Document object"""
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        extension = file_path_obj.suffix.lower()
        
        # Extract text based on file type
        if extension == '.pdf':
            text = DocumentProcessor.extract_text_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            text = DocumentProcessor.extract_text_from_docx(file_path)
        elif extension == '.txt':
            text = DocumentProcessor.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file format: {extension}")
            return None
        
        if not text or len(text.strip()) == 0:
            logger.warning(f"No text extracted from {file_path}")
            return None
        
        # Create Haystack Document
        doc = Document(
            content=text,
            meta={
                'file_path': str(file_path),
                'file_name': file_path_obj.name,
                'file_type': extension,
                'file_size': file_path_obj.stat().st_size
            }
        )
        
        return doc


class HaystackPipeline:
    """Main Haystack pipeline for RAG operations"""
    
    def __init__(
        self,
        model_name: str = "mistralai/Mistral-7B-Instruct-v0.2",
        embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
        device: str = "cpu",
        max_length: int = 512,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        use_hf_inference_api: bool = True,
        hf_api_key: str = "",
        use_cerebras_api: bool = False,
        cerebras_api_key: str = "",
        cerebras_model: str = "llama3.1-8b"
    ):
        self.model_name = model_name
        self.embedding_model = embedding_model
        self.device = device
        self.max_length = max_length
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_hf_inference_api = use_hf_inference_api
        self.hf_api_key = hf_api_key
        self.use_cerebras_api = use_cerebras_api
        self.cerebras_api_key = cerebras_api_key
        self.cerebras_model = cerebras_model
        
        # Initialize document store
        self.document_store = InMemoryDocumentStore()
        
        # Initialize components
        self._init_components()
        
        # Determine inference method
        if use_cerebras_api:
            inference_method = "Cerebras Inference API"
        elif use_hf_inference_api:
            inference_method = "Hugging Face Inference API"
        else:
            inference_method = "Local Model"
        
        logger.info(f"HaystackPipeline initialized with model: {model_name} using {inference_method}")
    
    def _init_components(self):
        """Initialize pipeline components"""
        # Resolve device to ComponentDevice object
        resolved_device = ComponentDevice.from_str(self.device) if isinstance(self.device, str) else self.device
        
        # Document preprocessors
        self.cleaner = DocumentCleaner()
        self.splitter = DocumentSplitter(
            split_by="word",
            split_length=self.chunk_size,
            split_overlap=self.chunk_overlap
        )
        
        # Embedders
        self.doc_embedder = SentenceTransformersDocumentEmbedder(
            model=self.embedding_model,
            device=resolved_device
        )
        self.text_embedder = SentenceTransformersTextEmbedder(
            model=self.embedding_model,
            device=resolved_device
        )
        
        # Retrievers
        self.bm25_retriever = InMemoryBM25Retriever(document_store=self.document_store)
        self.embedding_retriever = InMemoryEmbeddingRetriever(document_store=self.document_store)
        
        # Ranker for hybrid search
        self.ranker = TransformersSimilarityRanker(
            model="cross-encoder/ms-marco-MiniLM-L-6-v2"
        )
        
        # Document writer
        self.writer = DocumentWriter(document_store=self.document_store)
        
        logger.info("Pipeline components initialized")
    
    def warm_up(self):
        """Warm up all pipeline components"""
        try:
            logger.info("Warming up pipeline components...")
            
            # Warm up text embedder
            self.text_embedder.warm_up()
            
            # Warm up ranker
            self.ranker.warm_up()
            
            logger.info("Pipeline components warmed up successfully")
        except Exception as e:
            logger.error(f"Error warming up components: {str(e)}", exc_info=True)
            raise
    
    def ingest_document(self, file_path: str) -> bool:
        """Ingest a single document into the pipeline"""
        try:
            # Process document
            doc = DocumentProcessor.process_document(file_path)
            if not doc:
                return False
            
            # Create indexing pipeline with new component instances
            indexing_pipeline = Pipeline()
            
            # Create new instances to avoid component sharing issues
            cleaner = DocumentCleaner()
            splitter = DocumentSplitter(
                split_by="word",
                split_length=self.chunk_size,
                split_overlap=self.chunk_overlap
            )
            resolved_device = ComponentDevice.from_str(self.device) if isinstance(self.device, str) else self.device
            embedder = SentenceTransformersDocumentEmbedder(
                model=self.embedding_model,
                device=resolved_device
            )
            writer = DocumentWriter(document_store=self.document_store)
            
            indexing_pipeline.add_component("cleaner", cleaner)
            indexing_pipeline.add_component("splitter", splitter)
            indexing_pipeline.add_component("embedder", embedder)
            indexing_pipeline.add_component("writer", writer)
            
            indexing_pipeline.connect("cleaner", "splitter")
            indexing_pipeline.connect("splitter", "embedder")
            indexing_pipeline.connect("embedder", "writer")
            
            # Warm up the pipeline
            indexing_pipeline.warm_up()
            
            # Run pipeline
            result = indexing_pipeline.run({"cleaner": {"documents": [doc]}})
            
            logger.info(f"Successfully ingested document: {file_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error ingesting document {file_path}: {str(e)}", exc_info=True)
            return False
    
    def ingest_directory(self, directory_path: str) -> Dict[str, Any]:
        """Ingest all documents from a directory"""
        directory = Path(directory_path)
        
        if not directory.exists() or not directory.is_dir():
            logger.error(f"Directory not found: {directory_path}")
            return {"success": False, "error": "Directory not found"}
        
        results = {
            "success": True,
            "total": 0,
            "ingested": 0,
            "failed": 0,
            "files": []
        }
        
        # Supported extensions
        extensions = ['.pdf', '.docx', '.doc', '.txt']
        
        # Process all files
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                results["total"] += 1
                if self.ingest_document(str(file_path)):
                    results["ingested"] += 1
                    results["files"].append(str(file_path))
                else:
                    results["failed"] += 1
        
        logger.info(f"Ingested {results['ingested']}/{results['total']} documents from {directory_path}")
        return results
    
    def retrieve_bm25(self, query: str, top_k: int = 5) -> List[Document]:
        """Retrieve documents using BM25 (keyword search)"""
        try:
            results = self.bm25_retriever.run(query=query, top_k=top_k)
            return results.get("documents", [])
        except Exception as e:
            logger.error(f"Error in BM25 retrieval: {str(e)}")
            return []
    
    def retrieve_semantic(self, query: str, top_k: int = 5) -> List[Document]:
        """Retrieve documents using semantic search"""
        try:
            # Ensure text embedder is warmed up
            if not hasattr(self.text_embedder, '_warmed_up') or not self.text_embedder._warmed_up:
                self.text_embedder.warm_up()
            
            # Embed query
            query_embedding_result = self.text_embedder.run(text=query)
            query_embedding = query_embedding_result["embedding"]
            
            # Retrieve
            results = self.embedding_retriever.run(
                query_embedding=query_embedding,
                top_k=top_k
            )
            return results.get("documents", [])
        except Exception as e:
            logger.error(f"Error in semantic retrieval: {str(e)}")
            return []
    
    def retrieve_hybrid(self, query: str, top_k: int = 5) -> List[Document]:
        """Retrieve documents using hybrid search (BM25 + semantic)"""
        try:
            # Get results from both retrievers
            bm25_docs = self.retrieve_bm25(query, top_k=top_k * 2)
            semantic_docs = self.retrieve_semantic(query, top_k=top_k * 2)
            
            # Combine and deduplicate
            all_docs = bm25_docs + semantic_docs
            unique_docs = []
            seen_ids = set()
            
            for doc in all_docs:
                doc_id = doc.id
                if doc_id not in seen_ids:
                    unique_docs.append(doc)
                    seen_ids.add(doc_id)
            
            # Rank combined results
            if len(unique_docs) > 0:
                # Ensure ranker is warmed up
                if not hasattr(self.ranker, '_warmed_up') or not self.ranker._warmed_up:
                    self.ranker.warm_up()
                
                ranked_results = self.ranker.run(
                    query=query,
                    documents=unique_docs[:top_k * 2]
                )
                return ranked_results.get("documents", [])[:top_k]
            
            return []
            
        except Exception as e:
            logger.error(f"Error in hybrid retrieval: {str(e)}")
            return []
    
    def retrieve(
        self,
        query: str,
        top_k: int = 5,
        mode: str = "hybrid"
    ) -> List[Document]:
        """Retrieve documents using specified mode"""
        if mode == "bm25":
            return self.retrieve_bm25(query, top_k)
        elif mode == "semantic":
            return self.retrieve_semantic(query, top_k)
        else:  # hybrid
            return self.retrieve_hybrid(query, top_k)
    
    def generate_with_hf(
        self,
        prompt: str,
        max_new_tokens: int = 512,
        temperature: float = 0.7,
        do_sample: bool = True
    ) -> str:
        """Generate text using available API or local model"""
        if self.use_cerebras_api:
            return self._generate_with_cerebras(prompt, max_new_tokens, temperature)
        elif self.use_hf_inference_api:
            return self._generate_with_hf_api(prompt, max_new_tokens, temperature)
        else:
            return self._generate_with_local(prompt, max_new_tokens, temperature, do_sample)
    
    def _generate_with_cerebras(
        self,
        prompt: str,
        max_new_tokens: int = 512,
        temperature: float = 0.7
    ) -> str:
        """Generate text using Cerebras Inference API"""
        try:
            if not self.cerebras_api_key:
                raise ValueError("Cerebras API key not provided. Set CEREBRAS_API_KEY environment variable or pass it to the pipeline.")
            
            url = "https://api.cerebras.ai/v1/chat/completions"
            headers = {
                "Authorization": f"Bearer {self.cerebras_api_key}",
                "Content-Type": "application/json",
            }
            
            payload = {
                "model": self.cerebras_model,
                "messages": [
                    {"role": "user", "content": prompt}
                ],
                "max_tokens": max_new_tokens,
                "temperature": temperature,
            }
            
            logger.info(f"Calling Cerebras Inference API for model: {self.cerebras_model}")
            response = requests.post(url, headers=headers, json=payload, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                if "choices" in result and len(result["choices"]) > 0:
                    generated_text = result["choices"][0]["message"]["content"]
                    logger.info(f"Successfully generated {len(generated_text)} characters via Cerebras API")
                    return generated_text
                else:
                    logger.error(f"Unexpected Cerebras API response format: {result}")
                    return "Error: Unexpected response format from Cerebras API"
            elif response.status_code == 503:
                # Service unavailable
                error_msg = "Cerebras API is currently unavailable. Please try again later."
                logger.warning(error_msg)
                return error_msg
            else:
                error_msg = f"Cerebras API Error ({response.status_code}): {response.text}"
                logger.error(error_msg)
                return f"Error: {error_msg}"
            
        except requests.exceptions.Timeout:
            logger.error("Cerebras API request timed out")
            return "Error: Request timed out. Please try again."
        except Exception as e:
            logger.error(f"Error generating text with Cerebras API: {str(e)}", exc_info=True)
            return f"Error generating text: {str(e)}"
    
    def _generate_with_hf_api(
        self,
        prompt: str,
        max_new_tokens: int = 512,
        temperature: float = 0.7
    ) -> str:
        """Generate text using Hugging Face Inference API"""
        try:
            if not self.hf_api_key:
                raise ValueError("Hugging Face API key not provided. Set HF_API_KEY environment variable or pass it to the pipeline.")
            
            api_url = f"https://api-inference.huggingface.co/models/{self.model_name}"
            headers = {"Authorization": f"Bearer {self.hf_api_key}"}
            
            payload = {
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": max_new_tokens,
                    "temperature": temperature,
                    "return_full_text": False
                }
            }
            
            logger.info(f"Calling Hugging Face Inference API for model: {self.model_name}")
            response = requests.post(api_url, headers=headers, json=payload, timeout=120)
            
            if response.status_code == 200:
                result = response.json()
                if isinstance(result, list) and len(result) > 0:
                    generated_text = result[0].get("generated_text", "")
                    logger.info(f"Successfully generated {len(generated_text)} characters via API")
                    return generated_text
                else:
                    logger.error(f"Unexpected API response format: {result}")
                    return "Error: Unexpected response format from API"
            elif response.status_code == 503:
                # Model is loading
                error_msg = "Model is currently loading. Please try again in a few moments."
                logger.warning(error_msg)
                return error_msg
            else:
                error_msg = f"API Error ({response.status_code}): {response.text}"
                logger.error(error_msg)
                return f"Error: {error_msg}"
            
        except requests.exceptions.Timeout:
            logger.error("API request timed out")
            return "Error: Request timed out. The model might be loading, please try again."
        except Exception as e:
            logger.error(f"Error generating text with API: {str(e)}", exc_info=True)
            return f"Error generating text: {str(e)}"
    
    def _generate_with_local(
        self,
        prompt: str,
        max_new_tokens: int = 512,
        temperature: float = 0.7,
        do_sample: bool = True
    ) -> str:
        """Generate text using local Hugging Face model"""
        try:
            # Load tokenizer and model
            tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            model = AutoModelForCausalLM.from_pretrained(
                self.model_name,
                torch_dtype=torch.float16 if self.device == "cuda" else torch.float32,
                device_map=self.device,
                low_cpu_mem_usage=True
            )
            
            # Tokenize input
            inputs = tokenizer(prompt, return_tensors="pt").to(self.device)
            
            # Generate
            with torch.no_grad():
                outputs = model.generate(
                    **inputs,
                    max_new_tokens=max_new_tokens,
                    temperature=temperature,
                    do_sample=do_sample,
                    pad_token_id=tokenizer.eos_token_id
                )
            
            # Decode
            generated_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # Remove prompt from output
            if generated_text.startswith(prompt):
                generated_text = generated_text[len(prompt):].strip()
            
            return generated_text
            
        except Exception as e:
            logger.error(f"Error generating text with local model: {str(e)}", exc_info=True)
            return f"Error generating text: {str(e)}"
    
    def get_document_count(self) -> int:
        """Get total number of documents in store"""
        return self.document_store.count_documents()
    
    def clear_documents(self):
        """Clear all documents from store"""
        self.document_store.delete_documents(list(self.document_store.filter_documents().keys()))
        logger.info("Document store cleared")
